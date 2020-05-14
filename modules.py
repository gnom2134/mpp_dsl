from docx import Document
from bs4 import BeautifulSoup
from urllib import request, parse
from PIL import Image
from io import BytesIO


class HTMLObject:
    def __init__(self, html_page):
        self.content = BeautifulSoup(html_page, 'lxml')

    @staticmethod
    def dfs_text_contents(contents):
        line = ' '
        for j in contents:
            if hasattr(j, 'contents'):
                if j.get('class') is None or 'mwe-math-element' not in j.get('class'):
                    line += HTMLObject.dfs_text_contents(j.contents)
            else:
                line += str(j)
        return line

    @staticmethod
    def dfs_img_contents(contents):
        images = []
        for j in contents:
            if j.name == 'img':
                try:
                    image_src = j['src']
                    if not image_src.startswith('https:'):
                        image_src = 'https:' + image_src
                    file = BytesIO(request.urlopen(image_src).read())
                    img = Image.open(file)
                    images.append(img)
                except OSError as err:
                    print(f'OSError occurred {err}')
            if hasattr(j, 'contents'):
                res = HTMLObject.dfs_img_contents(j.contents)
                if res is not None:
                    images.extend(res)
        if len(images) > 0:
            return images
        else:
            return None

    def filter(self, filter_obj):
        tag = True
        args = {}
        if '@tag' in filter_obj:
            tag = filter_obj['@tag']
        if '@class' in filter_obj:
            args['class'] = filter_obj['@class']
        res = self.content.find_all(tag, args)
        res_result = []
        for i in res:
            line = self.dfs_text_contents(i.contents)
            images = self.dfs_img_contents(i.contents)
            res_result.append({'text': line, 'images': images})
        return res_result


class WebModule:
    def __init__(self, url):
        url = parse.quote(url, safe="%/:=&?~#+!$,;'@()*[]")
        page = request.urlopen(url)
        self.content = page.read().decode('utf8')
        page.close()

    def get(self):
        return HTMLObject(self.content)


class DocxModule:
    def __init__(self, filename):
        self.filename = filename
        self.document = Document()

    def write(self, data):
        text = data['text']
        if 'images' in data:
            images = data['images']
            for i in images:
                img = BytesIO()
                i.save(img, format='PNG')
                self.document.add_picture(img)
        pages = text.split('[new_page]')
        for i, v in enumerate(pages):
            self.document.add_paragraph(v)
            if i != len(pages) - 1:
                self.document.add_page_break()

    def __del__(self):
        self.document.save(self.filename)
